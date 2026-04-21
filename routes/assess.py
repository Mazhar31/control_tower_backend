import json
import os
import logging
import io
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import Optional, List
from database import get_db
from models import Assessment
from schemas import AssessmentSummary
from routes.deps import get_current_user
from leash_service import run_assessment, start_assessment, continue_assessment, finish_assessment
import httpx

logger = logging.getLogger(__name__)

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "..", "uploads")
ALLOWED_EXTS = {".pdf", ".docx"}

# ---------------------------------------------------------------------------
# Mock response for local dev fallback
# ---------------------------------------------------------------------------
USE_MOCK = os.getenv("USE_MOCK", "false").lower() == "true"

MOCK_RESPONSE = {
    "session_id": "mock-session-001",
    "company_name": "Test Company",
    "system_name": "Customer Analytics AI",
    "assessment_date": "2026-01-01",
    "risk_tier": "High",
    "frameworks_triggered": ["EU AI Act", "ISO 42001"],
    "critical_count": 2,
    "high_count": 3,
    "medium_count": 4,
    "low_count": 1,
    "gap_findings": [
        {
            "finding_id": "GAP-001",
            "category": "AI Governance",
            "title": "No AI Governance Framework",
            "description": "No formal AI governance policy or accountability structure exists.",
            "severity": "Critical",
            "current_state": "No governance framework, policy, or AI Officer identified.",
            "required_state": "EU AI Act Art. 17 requires documented risk management system.",
            "remediation": "Implement HAIG AI Governance Charter.",
            "remediation_template": "AI Governance/HAIG AI Governance Charter.docx",
            "framework_map": {
                "EU AI Act": ["Art. 9", "Art. 17"],
                "ISO 42001": ["5.1", "6.1"],
            },
        },
        {
            "finding_id": "GAP-002",
            "category": "Model Fairness",
            "title": "No Bias Testing Framework",
            "description": "No bias detection or testing process exists for AI models.",
            "severity": "High",
            "current_state": "No bias testing in place.",
            "required_state": "ISO 42001 requires fairness evaluation.",
            "remediation": "Implement bias testing pipeline.",
            "remediation_template": "Model Fairness/HAIG Bias Testing Template.docx",
            "framework_map": {
                "EU AI Act": ["Art. 10"],
                "ISO 42001": ["8.4"],
            },
        },
        {
            "finding_id": "GAP-003",
            "category": "Transparency",
            "title": "Missing Model Cards",
            "description": "No model cards exist for deployed AI models.",
            "severity": "Medium",
            "current_state": "No documentation.",
            "required_state": "ISO 42001 requires model documentation.",
            "remediation": "Create model cards.",
            "remediation_template": "Documentation/HAIG Model Card Template.docx",
            "framework_map": {
                "ISO 42001": ["8.6"],
                "EU AI Act": ["Art. 13"],
            },
        },
    ],
    "framework_coverage": {
        "EU AI Act": {"gaps": 3, "pct_compliant": 42},
        "ISO 42001": {"gaps": 3, "pct_compliant": 55},
    },
    "shelby_controls": {
        "has_ai_usage_policy": False,
        "employee_ai_training": False,
        "ai_output_review_process": True,
        "dependency_risk_assessed": False,
    },
    "report_url": "https://example.com/mock-report.pdf",
}

router = APIRouter(prefix="/assess", tags=["assess"])


@router.post("")
async def submit_assessment(
    companyName: str = Form(...),
    systemName: str = Form(...),
    systemDescription: str = Form(...),
    aihcsResponse: str = Form(...),
    geography: str = Form(""),
    usStates: str = Form(""),
    industry: str = Form(""),
    sectorRegs: str = Form(""),
    selectedFrameworks: str = Form(""),
    aiType: str = Form(""),
    decisionImpact: str = Form(""),
    existingDocs: str = Form(""),
    dataTypes: str = Form(""),
    infrastructure: str = Form(""),
    scaleEstimate: str = Form(""),
    aihcsDetail: str = Form(""),
    deploymentStage: str = Form("production"),
    additionalContext: str = Form(""),
    contactName: str = Form(""),
    contactTitle: str = Form(""),
    contactEmail: str = Form(""),
    nextAuditDate: str = Form(""),
    files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
    user=Depends(get_current_user),
):
    user_id = int(user["sub"])

    # Handle optional file uploads
    uploaded_file_names = []
    all_file_texts = []
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    for upload in files:
        if not upload.filename:
            continue
        ext = os.path.splitext(upload.filename)[1].lower()
        if ext not in ALLOWED_EXTS:
            raise HTTPException(status_code=400, detail=f"{upload.filename}: only PDF and DOCX files are accepted")
        file_bytes = await upload.read()
        dest = os.path.join(UPLOAD_DIR, f"user{user_id}_{upload.filename}")
        with open(dest, "wb") as f:
            f.write(file_bytes)
        uploaded_file_names.append(upload.filename)
        try:
            if ext == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            elif ext == ".docx":
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                text = ""
            if text:
                all_file_texts.append(f"--- Document: {upload.filename} ---\n{text[:8000]}")
        except Exception as e:
            logger.warning("[FILE] Could not extract text from %s: %s", upload.filename, e)

    # Build intake payload exactly matching Leash API spec
    full_additional_context = (additionalContext + ("\n\n" + "\n\n".join(all_file_texts)) if all_file_texts else additionalContext).strip()

    intake_data = {
        "companyName": companyName,
        "systemName": systemName,
        "systemDescription": systemDescription,
        "industry": industry,
        "geography": geography,
        "usStates": [s.strip() for s in usStates.split(",") if s.strip()],
        "selectedFrameworks": ["EU AI Act", "ISO 42001", "HIPAA", "NIST AI RMF"],
        "aihcsResponse": aihcsResponse,
        "deploymentStage": deploymentStage if deploymentStage else "production",
        "dataTypes": dataTypes,
        "additionalContext": full_additional_context,
    }

    # Call The Leash (or mock)
    if USE_MOCK:
        logger.info("[MOCK] Returning mock response — set USE_MOCK=false to call real API")
        result = MOCK_RESPONSE
    else:
        logger.info("[LEASH] Calling real API for user=%s company=%s system=%s", user_id, companyName, systemName)
        logger.info("[LEASH] Payload: %s", json.dumps(intake_data, indent=2))
        try:
            raw = await run_assessment(intake_data)
            # Real response wraps findings inside taiga_token — unwrap and merge
            taiga_token = raw.get("taiga_token", {})
            result = {
                **taiga_token,
                "report_url": raw.get("report_url", ""),
                "report_s3_key": raw.get("report_s3_key", ""),
                "total_findings": raw.get("total_findings", len(taiga_token.get("gap_findings", []))),
                "audit_summary": raw.get("audit_summary", {}),
            }
            logger.info("[LEASH] Success — session_id=%s risk_tier=%s findings=%s",
                result.get("session_id"), result.get("risk_tier"),
                len(result.get("gap_findings", [])))
        except httpx.TimeoutException:
            logger.error("[LEASH] Timed out after %ss", 900)
            raise HTTPException(status_code=504, detail="Assessment timed out. Please try again.")
        except Exception as e:
            logger.error("[LEASH] Failed: %s", str(e))
            raise HTTPException(status_code=502, detail=f"Assessment service error: {str(e)}")

    # Calculate prior score from user's last assessment for trend
    prior = (
        db.query(Assessment)
        .filter(Assessment.user_id == user_id)
        .order_by(Assessment.id.desc())
        .first()
    )
    if prior and prior.response_json:
        try:
            prior_result = json.loads(prior.response_json)
            prior_coverage = prior_result.get("framework_coverage", {})
            vals = [v["pct_compliant"] for v in prior_coverage.values() if isinstance(v.get("pct_compliant"), (int, float))]
            result["prior_score"] = round(sum(vals) / len(vals)) if vals else None
        except Exception:
            result["prior_score"] = None
    else:
        result["prior_score"] = None

    if nextAuditDate:
        result["next_audit_date"] = nextAuditDate

    # Persist to DB
    assessment = Assessment(
        user_id=user_id,
        session_id=result.get("session_id"),
        company_name=result.get("company_name"),
        system_name=result.get("system_name"),
        assessment_date=result.get("assessment_date"),
        risk_tier=result.get("risk_tier"),
        response_json=json.dumps(result),
        uploaded_file_name=json.dumps(uploaded_file_names) if uploaded_file_names else None,
        uploaded_file_path=None,
    )
    db.add(assessment)
    db.commit()
    db.refresh(assessment)

    return {"assessment_id": assessment.id, "result": result}


def _build_prior_score(db, user_id):
    prior = (
        db.query(Assessment)
        .filter(Assessment.user_id == user_id)
        .order_by(Assessment.id.desc())
        .first()
    )
    if prior and prior.response_json:
        try:
            prior_coverage = json.loads(prior.response_json).get("framework_coverage", {})
            vals = [v["pct_compliant"] for v in prior_coverage.values() if isinstance(v.get("pct_compliant"), (int, float))]
            return round(sum(vals) / len(vals)) if vals else None
        except Exception:
            return None
    return None


def _save_assessment(db, user_id, result, uploaded_file_names, nextAuditDate):
    result["prior_score"] = _build_prior_score(db, user_id)
    if nextAuditDate:
        result["next_audit_date"] = nextAuditDate
    assessment = Assessment(
        user_id=user_id,
        session_id=result.get("session_id"),
        company_name=result.get("company_name"),
        system_name=result.get("system_name"),
        assessment_date=result.get("assessment_date"),
        risk_tier=result.get("risk_tier"),
        response_json=json.dumps(result),
        uploaded_file_name=json.dumps(uploaded_file_names) if uploaded_file_names else None,
        uploaded_file_path=None,
    )
    db.add(assessment)
    db.commit()
    db.refresh(assessment)
    return assessment


@router.post("/start")
async def start_assessment_route(
    companyName: str = Form(...),
    systemName: str = Form(...),
    systemDescription: str = Form(...),
    aihcsResponse: str = Form(...),
    geography: str = Form(""),
    usStates: str = Form(""),
    industry: str = Form(""),
    sectorRegs: str = Form(""),
    selectedFrameworks: str = Form(""),
    aiType: str = Form(""),
    decisionImpact: str = Form(""),
    existingDocs: str = Form(""),
    dataTypes: str = Form(""),
    infrastructure: str = Form(""),
    scaleEstimate: str = Form(""),
    aihcsDetail: str = Form(""),
    deploymentStage: str = Form("production"),
    additionalContext: str = Form(""),
    contactName: str = Form(""),
    contactTitle: str = Form(""),
    contactEmail: str = Form(""),
    nextAuditDate: str = Form(""),
    files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
    user=Depends(get_current_user),
):
    """Phase 1: submit intake, get back questions or go straight to result."""
    user_id = int(user["sub"])

    uploaded_file_names = []
    all_file_texts = []
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    for upload in files:
        if not upload.filename:
            continue
        ext = os.path.splitext(upload.filename)[1].lower()
        if ext not in ALLOWED_EXTS:
            raise HTTPException(status_code=400, detail=f"{upload.filename}: only PDF and DOCX files are accepted")
        file_bytes = await upload.read()
        with open(os.path.join(UPLOAD_DIR, f"user{user_id}_{upload.filename}"), "wb") as f:
            f.write(file_bytes)
        uploaded_file_names.append(upload.filename)
        try:
            if ext == ".pdf":
                from pypdf import PdfReader
                text = "\n".join(p.extract_text() or "" for p in PdfReader(io.BytesIO(file_bytes)).pages)
            elif ext == ".docx":
                from docx import Document
                text = "\n".join(p.text for p in Document(io.BytesIO(file_bytes)).paragraphs if p.text.strip())
            else:
                text = ""
            if text:
                all_file_texts.append(f"--- Document: {upload.filename} ---\n{text[:8000]}")
        except Exception as e:
            logger.warning("[FILE] Could not extract text from %s: %s", upload.filename, e)

    full_additional_context = (additionalContext + ("\n\n" + "\n\n".join(all_file_texts)) if all_file_texts else additionalContext).strip()

    intake_data = {
        "companyName": companyName,
        "systemName": systemName,
        "systemDescription": systemDescription,
        "industry": industry,
        "geography": geography,
        "usStates": [s.strip() for s in usStates.split(",") if s.strip()],
        "selectedFrameworks": ["EU AI Act", "ISO 42001", "HIPAA", "NIST AI RMF"],
        "aihcsResponse": aihcsResponse,
        "deploymentStage": deploymentStage if deploymentStage else "production",
        "dataTypes": dataTypes,
        "additionalContext": full_additional_context,
    }

    if USE_MOCK:
        # Mock: skip questions, return final result directly
        result = MOCK_RESPONSE
        assessment = _save_assessment(db, user_id, result, uploaded_file_names, nextAuditDate)
        return {"status": "complete", "assessment_id": assessment.id, "result": result}

    try:
        leash_result = await start_assessment(intake_data)
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Assessment timed out. Please try again.")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Assessment service error: {str(e)}")

    if leash_result.get("action_required") == "answer_questions":
        # Return questions to the frontend — session_id needed for /continue
        return {
            "status": "questions",
            "session_id": leash_result["session_id"],
            "questions": leash_result.get("questions", []),
            "round": leash_result.get("round", 1),
            "_meta": {"uploaded_file_names": uploaded_file_names, "nextAuditDate": nextAuditDate},
        }

    # No questions — run audit+generate immediately
    try:
        raw = await finish_assessment(leash_result["session_id"])
        taiga_token = raw.get("taiga_token", {})
        result = {
            **taiga_token,
            "report_url": raw.get("report_url", ""),
            "report_s3_key": raw.get("report_s3_key", ""),
            "total_findings": raw.get("total_findings", len(taiga_token.get("gap_findings", []))),
            "audit_summary": raw.get("audit_summary", {}),
        }
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Assessment timed out. Please try again.")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Assessment service error: {str(e)}")

    assessment = _save_assessment(db, user_id, result, uploaded_file_names, nextAuditDate)
    return {"status": "complete", "assessment_id": assessment.id, "result": result}


@router.post("/continue")
async def continue_assessment_route(
    session_id: str = Form(...),
    answers: str = Form(...),          # JSON string: [{question_index, answer}, ...]
    uploaded_file_names: str = Form(""),  # JSON string
    nextAuditDate: str = Form(""),
    db: Session = Depends(get_db),
    user=Depends(get_current_user),
):
    """Phase 2: submit answers, get back more questions or final result."""
    user_id = int(user["sub"])

    try:
        parsed_answers = json.loads(answers)
        parsed_file_names = json.loads(uploaded_file_names) if uploaded_file_names else []
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid answers format.")

    try:
        raw = await continue_assessment(session_id, parsed_answers)
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Assessment timed out. Please try again.")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Assessment service error: {str(e)}")

    # Check if more questions needed
    if raw.get("action_required") == "answer_questions":
        return {
            "status": "questions",
            "session_id": raw["session_id"],
            "questions": raw.get("questions", []),
            "round": raw.get("round", 1),
            "_meta": {"uploaded_file_names": parsed_file_names, "nextAuditDate": nextAuditDate},
        }

    # No more questions — run audit + generate
    try:
        raw = await finish_assessment(raw.get("session_id", session_id))
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Assessment timed out. Please try again.")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Assessment service error: {str(e)}")

    taiga_token = raw.get("taiga_token", {})
    result = {
        **taiga_token,
        "report_url": raw.get("report_url", ""),
        "report_s3_key": raw.get("report_s3_key", ""),
        "total_findings": raw.get("total_findings", len(taiga_token.get("gap_findings", []))),
        "audit_summary": raw.get("audit_summary", {}),
    }

    assessment = _save_assessment(db, user_id, result, parsed_file_names, nextAuditDate)
    return {"status": "complete", "assessment_id": assessment.id, "result": result}


@router.get("/history")
def get_history(db: Session = Depends(get_db), user=Depends(get_current_user)):
    rows = (
        db.query(Assessment)
        .filter(Assessment.user_id == int(user["sub"]))
        .order_by(Assessment.id.desc())
        .all()
    )
    return [
        {
            "id": r.id,
            "company_name": r.company_name,
            "system_name": r.system_name,
            "assessment_date": r.assessment_date,
            "risk_tier": r.risk_tier,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in rows
    ]


@router.get("/{assessment_id}")
def get_assessment(assessment_id: int, db: Session = Depends(get_db), user=Depends(get_current_user)):
    row = db.query(Assessment).filter(
        Assessment.id == assessment_id,
        Assessment.user_id == int(user["sub"]),
    ).first()
    if not row:
        raise HTTPException(status_code=404, detail="Assessment not found")
    return {"assessment_id": row.id, "result": json.loads(row.response_json)}
